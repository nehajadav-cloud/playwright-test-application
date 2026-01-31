import argparse
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont


TEST_CASES = [
    {
        "id": 1,
        "title": "Start and stop the application and verify",
        "steps": [
            "Start server on a test port",
            "Open the home page and verify UI loads",
            "Stop the server and verify it no longer responds",
        ],
        "expected": "App starts and stops cleanly; UI loads while running."
    },
    {
        "id": 2,
        "title": "Login successfully with admin/admin123",
        "steps": [
            "Open login page",
            "Enter admin credentials",
            "Click Sign in and verify redirect to /employees",
        ],
        "expected": "User is logged in and sees Employees page."
    },
    {
        "id": 3,
        "title": "Add testing user after verifying it doesn't exist",
        "steps": [
            "Search for a unique test ID",
            "Confirm no matching rows",
            "Fill form and click Save",
            "Verify the new user appears",
        ],
        "expected": "Employee is created and visible in table."
    },
    {
        "id": 4,
        "title": "Delete employee works (remove testing user)",
        "steps": [
            "Create a temporary employee",
            "Delete the employee",
            "Verify it no longer appears",
        ],
        "expected": "Employee is deleted."
    },
    {
        "id": 5,
        "title": "Edit user works (add, edit, save, verify)",
        "steps": [
            "Create a temporary employee",
            "Edit the name and save",
            "Verify updated name appears",
        ],
        "expected": "Employee updates are saved."
    },
    {
        "id": 6,
        "title": "Delete employee works again (cleanup another user)",
        "steps": [
            "Create a second temporary employee",
            "Delete the employee",
            "Verify it no longer appears",
        ],
        "expected": "Second employee is deleted."
    },
]


def slugify(text: str) -> str:
    import re
    return re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")

def normalize_title(text: str) -> str:
    import re
    return re.sub(r"^\d+\)\s*", "", text)

def normalize_status(status: str) -> str:
    if status == "unexpected":
        return "failed"
    return status

def annotate_image(src: Path, title: str, browser: str) -> Path:
    img = Image.open(src).convert("RGB")
    banner_height = 48
    new_img = Image.new("RGB", (img.width, img.height + banner_height), (20, 20, 20))
    new_img.paste(img, (0, banner_height))

    draw = ImageDraw.Draw(new_img)
    text = f"{title} | {browser}"
    try:
        font = ImageFont.truetype("arial.ttf", 18)
    except Exception:
        font = ImageFont.load_default()
    draw.text((12, 14), text, font=font, fill=(255, 255, 255))

    out = src.with_name(f"annotated-{src.name}")
    new_img.save(out)
    return out

def read_json(path: Path):
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def extract_results(report):
    tests = []

    def walk_suite(suite):
        for spec in suite.get("specs", []):
            for t in spec.get("tests", []):
                for r in t.get("results", []):
                    tests.append({
                        "title": spec.get("title", ""),
                        "status": r.get("status", "unknown"),
                        "duration": r.get("duration", 0),
                        "project": t.get("projectName", "unknown"),
                        "startTime": r.get("startTime", ""),
                        "error": (r.get("error") or {}).get("message", ""),
                    })
        for child in suite.get("suites", []):
            walk_suite(child)

    for suite in report.get("suites", []):
        walk_suite(suite)

    return tests


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--json", required=True, action="append")
    parser.add_argument("--run-id", required=True)
    parser.add_argument("--out", required=True)
    args = parser.parse_args()

    results = []
    for json_path in args.json:
        report = read_json(Path(json_path))
        results.extend(extract_results(report))
    result_map = {}
    for r in results:
        key = normalize_title(r["title"])
        result_map.setdefault(key, []).append(r)
    summary_rows = []

    doc = Document()
    doc.add_heading("Playwright Test Execution Report", 0)
    doc.add_paragraph(f"Run ID: {args.run_id}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    shots_dir = Path("test-results") / "screenshots" / args.run_id

    for tc in TEST_CASES:
        title = f"{tc['id']}. {tc['title']}"
        doc.add_heading(title, level=1)
        doc.add_paragraph("Description:")
        doc.add_paragraph(tc["title"])
        doc.add_paragraph("Steps:")
        for step in tc["steps"]:
            doc.add_paragraph(step, style="List Number")
        doc.add_paragraph(f"Expected: {tc['expected']}")

        results_for_test = result_map.get(tc["title"], [])
        if not results_for_test:
            doc.add_paragraph("Status: unknown")
        else:
            doc.add_paragraph("Status by browser:")
            for r in results_for_test:
                status = normalize_status(r["status"])
                doc.add_paragraph(
                    f"{r['project']}: {status} (ms: {r['duration']})",
                    style="List Bullet"
                )
                summary_rows.append({
                    "id": tc["id"],
                    "title": tc["title"],
                    "browser": r["project"],
                    "status": status,
                    "startTime": r.get("startTime", ""),
                    "duration": r.get("duration", 0),
                    "error": (r.get("error") or "").splitlines()[0] if r.get("error") else "",
                })

        slug = slugify(tc["title"])
        if shots_dir.exists():
            matches = sorted(shots_dir.glob(f"{slug}-*.png"))
            for img in matches:
                browser = "browser"
                parts = img.stem.split("-")
                if len(parts) >= 2:
                    browser = parts[-2]
                annotated = annotate_image(img, tc["title"], browser)
                doc.add_paragraph(f"Snapshot: {img.name}")
                doc.add_picture(str(annotated), width=Inches(5.5))

        if tc["id"] == 1:
            log_path = shots_dir / "start-stop-log.json"
            if log_path.exists():
                log_data = json.loads(log_path.read_text(encoding="utf-8"))
                doc.add_paragraph(f"Start time: {log_data.get('startedAt', '')}")
                doc.add_paragraph(f"Stop time: {log_data.get('stoppedAt', '')}")
                logs = log_data.get("logs", [])
                if logs:
                    doc.add_paragraph("Server output:")
                    for line in logs:
                        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Summary Table", level=1)
    table = doc.add_table(rows=1, cols=7)
    hdr = table.rows[0].cells
    hdr[0].text = "TC"
    hdr[1].text = "Test Name"
    hdr[2].text = "Browser"
    hdr[3].text = "Status"
    hdr[4].text = "Start Time"
    hdr[5].text = "Duration (ms)"
    hdr[6].text = "Failure Reason"

    for row in summary_rows:
        cells = table.add_row().cells
        cells[0].text = str(row["id"])
        cells[1].text = row["title"]
        cells[2].text = row["browser"]
        cells[3].text = row["status"]
        cells[4].text = row["startTime"]
        cells[5].text = str(row["duration"])
        cells[6].text = row["error"]

    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
